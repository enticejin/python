from os import path, remove
from tkinter import *
from tkinter import messagebox, filedialog
from win32com import client
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document


def pdf2docx(pdf_name, docx_name):
    try:
        if path.exists(docx_name):
            remove(docx_name)
        if path.exists(pdf_name):
            # rb以二进制读模式打开本地pdf文件
            fn = open(pdf_name, 'rb')
            # 创建一个pdf文档分析器
            parser = PDFParser(fn)
            # 创建一个PDF文档
            doc = PDFDocument()
            # 连接分析器 与文档对象
            parser.set_document(doc)
            doc.set_parser(parser)
            # 提供初始化密码doc.initialize("lianxipython")
            # 如果没有密码 就创建一个空的字符串
            doc.initialize("")
            # 检测文档是否提供txt转换，不提供就忽略
            if not doc.is_extractable:
                raise PDFTextExtractionNotAllowed
            else:
                # 创建PDf资源管理器
                resource = PDFResourceManager()
                # 创建一个PDF参数分析器
                laparams = LAParams()
                # 创建聚合器,用于读取文档的对象
                device = PDFPageAggregator(resource, laparams=laparams)
                # 创建解释器，对文档编码，解释成Python能够识别的格式
                interpreter = PDFPageInterpreter(resource, device)
                # 循环遍历列表，每次处理一页的内容
                # doc.get_pages() 获取page列表
                for page in doc.get_pages():
                    # 利用解释器的process_page()方法解析读取单独页数
                    interpreter.process_page(page)
                    # 使用聚合器get_result()方法获取内容
                    layout = device.get_result()
                    # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
                    for out in layout:
                        # 判断是否含有get_text()方法，获取我们想要的文字
                        if hasattr(out, "get_text"):
                            # print(out.get_text(), type(out.get_text()))
                            content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                            # with open('test.txt','a') as f:
                            #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                            document.add_paragraph(content)  # 添加段落，样式为unordered list类型
        else:
            messagebox.showinfo(title='提示', message='文件不存在')
    except:
        messagebox.showinfo(title='提示', message='未知原因导致转换失败')


def openfile():
    file = filedialog.askopenfilename(title="打开文件", filetypes=[('All Files', '*.pdf')])
    filespath_text.set(file)


def on_click():
    filepath = filespath_text.get()
    if filepath[-4:] == '.pdf':
        docx_name = filepath.replace('.pdf', '.docx')
        pdf_name = filepath.replace("\\", "/")
        pdf2docx(pdf_name, docx_name)
        document.save(docx_name)  # 保存这个文档
    else:
        messagebox.showinfo(title='提示', message='文件不存在或类型错误(*.pdf)')


if __name__ == '__main__':
    # UI
    root = Tk()
    root.title("pdf2word")
    root.geometry('265x66')
    root.resizable(width=False, height=False)

    filespathL = Label(root, text="path:", font=11)
    filespathL.grid(row=0, column=0, sticky=E)
    filespath_text = StringVar()
    filespathE = Entry(root, textvariable=filespath_text, font=11)
    filespath_text.set("")
    filespathE.grid(row=0, column=1, sticky=E)

    Button(root, bd=5, text="open", font=11, command=openfile).grid(row=1, column=0, sticky=E)
    Button(root, bd=5, text="pdf2word", font=11, command=on_click).grid(row=1, column=1, sticky=E)

    document = Document()
    root.mainloop()
